"""職務経歴書 2026年5月版を 2023-07-20 版を base に生成"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

OUTPUT = '/tmp/doda-docs/shokumu-2026-0502-v2.docx'

doc = Document()

# Set default font (日本語フォント MS Pゴシック)
style = doc.styles['Normal']
style.font.name = 'MS Pゴシック'
style.font.size = Pt(10.5)

# === Title (centered) ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('職務経歴書')
run.bold = True
run.font.size = Pt(14)

# === Name (right-aligned) ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run('氏名：仲 啓輔')

# === 職務経歴概要 ===
p = doc.add_paragraph()
p.add_run('■職務経歴概要').bold = True

summary_lines = [
    '大学を卒業後、西谷憩いの家にて約3年間介護福祉士として勤務しておりました。',
    '上記後、ダイハツ工業株式会社、株式会社フルタイムシステムにてオペレーター業務に従事しておりました。',
    'その後、自身にスキルをつけるため、HAL大阪に入学し学び直しをいたしました。',
    '在学中にトランスコスモス株式会社に中途採用いただいたため中退し、約12年間データセンターのオペレーター業務（システム運用・ヘルプデスク・テクニカルサポート）に従事しておりました。',
    '2020年9月にトランスコスモスを退職後、個人事業主として配送業務（ワタミの宅食）等に従事し、2023年7月以降は家庭の事情により主夫として家事を担いながら、AIエージェント運用・OSS開発・技術コミュニティ活動（Microsoft 365 勉強会主催等）を継続しております。',
    '2024年4月以降は、株式会社グロース・コンティニュー様より業務委託契約にて月次でPCキッティング業務（Windows10/11、複数台/月）を継続的に受託しており、トランスコスモス時代のキッティング経験を活かして稼働中です。',
]
for line in summary_lines:
    doc.add_paragraph(line)

# === 会社履歴 ===
doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('■会社履歴').bold = True

table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
hdr = table.rows[0].cells
hdr[0].text = '期間'
hdr[1].text = '会社名'

companies = [
    ('2000年10月〜2003年09月', '西谷憩いの家（正社員にて勤務）'),
    ('2003年10月〜2004年09月', 'ダイハツ工業株式会社（派遣社員にて勤務）'),
    ('2004年10月〜2005年03月', '株式会社フルタイムシステム（契約社員にて勤務）'),
    ('2007年10月〜2020年09月', 'トランスコスモス株式会社(正社員にて勤務）'),
    ('2022年03月〜2022年12月', '個人事業主'),
    ('2023年07月〜現在', '個人事業主・主夫'),
    ('2024年04月〜現在', '株式会社グロース・コンティニュー（業務委託・PCキッティング）'),
]
for period, name in companies:
    row = table.add_row().cells
    row[0].text = period
    row[1].text = name

# === 職務経歴詳細 ===
doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('■職務経歴詳細').bold = True


def add_job_section(title, period, duty, detail, special):
    """1 つの会社の詳細テーブルを追加"""
    p = doc.add_paragraph()
    p.add_run(f'{title}   {period}').bold = True

    table = doc.add_table(rows=2, cols=3)
    table.style = 'Table Grid'

    # Header
    hdr = table.rows[0].cells
    hdr[0].text = '期間'
    hdr[1].text = '業務内容'
    hdr[2].text = '特記事項'

    # Body
    body = table.rows[1].cells
    body[0].text = period.replace('〜', '\n～')
    body[1].text = f'【職務】\n{duty}\n【職務詳細】\n{detail}'
    body[2].text = special
    doc.add_paragraph('')


add_job_section(
    title='西谷憩いの家（正社員にて勤務）',
    period='2000年10月〜2003年09月',
    duty='介護福祉士',
    detail='介護の現場です。',
    special='夜勤経験あり',
)

add_job_section(
    title='ダイハツ工業株式会社（派遣社員にて勤務）',
    period='2003年10月〜2004年09月',
    duty='オペレーター',
    detail='印刷物の仕分け、運搬等',
    special='―',
)

add_job_section(
    title='株式会社フルタイムシステム（契約社員にて勤務）',
    period='2004年10月〜2005年03月',
    duty='オペレーター',
    detail='コールセンターでの電話対応等',
    special='―',
)

# トランスコスモス - special table with extra column for env
p = doc.add_paragraph()
p.add_run('トランスコスモス株式会社（正社員にて勤務）   2007年10月〜2020年09月').bold = True

t = doc.add_table(rows=2, cols=3)
t.style = 'Table Grid'
hdr = t.rows[0].cells
hdr[0].text = '期間'
hdr[1].text = '業務内容'
hdr[2].text = '特記事項・開発環境'
body = t.rows[1].cells
body[0].text = '2007年10月\n～2020年09月'
body[1].text = (
    '【職務】\n'
    'オペレーター（システム運用）\n'
    '【職務詳細】\n'
    'データセンターでのオペレーター業務。\n'
    'ヘルプデスク、テクニカルサポート\n'
    'システムの技術的な問合せに対する顧客対応\n'
    'ネットワーク・サーバーの運用・保守'
)
body[2].text = (
    '【環境】Linux\n'
    '・介護休職を計1年3ヶ月取得\n'
    '・夜勤経験あり\n\n'
    '［経験のある主な配属先］\n'
    '・住生情報システム株式会社\n'
    '・株式会社QTnet\n'
    '・株式会社メディセオ（医薬品卸）\n'
    '・住友電工工業株式会社'
)
doc.add_paragraph('')

add_job_section(
    title='個人事業主',
    period='2022年03月〜2022年12月',
    duty='ワタミの宅食',
    detail='宅食の配送',
    special='―',
)

# === NEW: 個人事業主・主夫 (2023年07月〜現在) ===
p = doc.add_paragraph()
p.add_run('個人事業主・主夫（AI/OSS活動）   2023年07月〜現在').bold = True

t = doc.add_table(rows=2, cols=3)
t.style = 'Table Grid'
hdr = t.rows[0].cells
hdr[0].text = '期間'
hdr[1].text = '業務内容'
hdr[2].text = '特記事項・環境'
body = t.rows[1].cells
body[0].text = '2023年07月\n～現在'
body[1].text = (
    '【職務】\n'
    '個人事業主・主夫\n'
    '【職務詳細】\n'
    '・家事業務（家事全般）\n'
    '・AIエージェント運用\n'
    '  （Anthropic Claude Code, OpenAI Codex CLI）\n'
    '・OSS開発・公開\n'
    '  （GitHub: wirelessml/pdf-reader、shibu-video-editor 等）\n'
    '・技術記事執筆（Substack 連載・仲啓輔名義）\n'
    '・技術コミュニティ運営\n'
    '  （Microsoft 365 勉強会主催・第53回 2026年4月開催 等）'
)
body[2].text = (
    '【環境】\n'
    'macOS / Windows 11 / Linux\n'
    'Python / PowerShell / Bash\n'
    'SSH / Tailscale\n'
    'Git / GitHub\n\n'
    '【取得資格】\n'
    '・LPIC1（保有継続）'
)
doc.add_paragraph('')

# === NEW v2: 株式会社グロース・コンティニュー（業務委託） ===
p = doc.add_paragraph()
p.add_run('株式会社グロース・コンティニュー（業務委託）   2024年04月〜現在').bold = True

t2 = doc.add_table(rows=2, cols=3)
t2.style = 'Table Grid'
hdr = t2.rows[0].cells
hdr[0].text = '期間'
hdr[1].text = '業務内容'
hdr[2].text = '特記事項・実績'
body = t2.rows[1].cells
body[0].text = '2024年04月\n～現在'
body[1].text = (
    '【職務】\n'
    'PCキッティング業務委託\n'
    '【職務詳細】\n'
    '株式会社グロース・コンティニュー様より業務委託契約にて、\n'
    '同社入社者向けPCキッティングを月次で受託。\n'
    '・Windows10 / Windows11 のクリーンセットアップ\n'
    '・業務アプリケーションのインストール・設定\n'
    '・Active Directory 参加・GPO 適用\n'
    '・初期設定確認・引渡し前検証\n'
    '・複数台/月（最大2-数台）を継続的に対応'
)
body[2].text = (
    '・月次継続契約（源泉徴収対象）\n'
    '・在宅で対応可能な体制を構築\n'
    '・トランスコスモス時代の住友電工配属で\n'
    '  全社員分Windows10セットアップ経験を活用\n'
    '・1年以上の継続実績'
)
doc.add_paragraph('')

# === PC スキル ===
p = doc.add_paragraph()
p.add_run('■PCスキル').bold = True

skills = [
    '・LPIC1',
    '・Linuxの知識、経験',
    '・住友電工工業株式会社に配属された際キッティングを経験。全社員分のWindows10パソコンを設定。',
    '・AIコーディングエージェント運用（Anthropic Claude Code, OpenAI Codex CLI）',
    '・Python / PowerShell / Bash による自動化スクリプト開発',
    '・macOS / Windows 11 / Linux マルチプラットフォーム運用',
    '・SSH / Tailscale を用いたリモート管理',
    '・Git / GitHub によるOSSプロジェクト運用',
]
for s in skills:
    doc.add_paragraph(s)

# === 自己 PR ===
doc.add_paragraph('')
p = doc.add_paragraph()
p.add_run('■自己PR').bold = True

pr_sections = [
    (
        '業界横断の対応力',
        '医療現場（介護福祉士3年）、コールセンター（電話対応）、'
        'データセンター（13年・ヘルプデスク・テクニカルサポート・システム運用）と'
        '業界横断で就業した経験から、エンドユーザーの業務を理解した上での'
        'IT支援ができることが強みです。とくに病院・介護施設での実務経験は、'
        '医療従事者の方とのコミュニケーションを円滑にする土台になっております。',
    ),
    (
        '業務効率化',
        'スピード感を持ち業務するため、作業の効率化を常に意識しております。'
        '近年はAIエージェントを活用した業務自動化にも積極的に取り組んでおります。',
    ),
    (
        '学習意欲',
        '常にエンドユーザーからの問合せに対応できるよう最新の情報を把握し、'
        '幅広く知識を取得しております。技術コミュニティ（Microsoft 365 勉強会）'
        'の主催を通じて、自身の学習成果を共有することにも取り組んでおります。',
    ),
    (
        '自作PC作成',
        '趣味はパソコンの自作や将棋です。ハードウェアの知識を業務でも活かせます。',
    ),
]

for title, body in pr_sections:
    p = doc.add_paragraph()
    p.add_run(title).bold = True
    doc.add_paragraph(body)
    doc.add_paragraph('')

# === 以上 ===
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
p.add_run('以上')

doc.save(OUTPUT)
print(f"✅ Saved: {OUTPUT}")

# Verify
import os
size = os.path.getsize(OUTPUT)
print(f"File size: {size} bytes ({size/1024:.1f} KB)")

# Re-read to verify
doc2 = Document(OUTPUT)
print(f"Paragraphs: {len(doc2.paragraphs)}")
print(f"Tables: {len(doc2.tables)}")
